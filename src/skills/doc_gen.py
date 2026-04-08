import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.state import BrandContext

def generate_docx(content: str, brand: BrandContext, output_path: str = "output.docx") -> str:
    """
    Generates a brand-compliant DOCX file.
    """
    doc = Document()

    # 1. Title/Header
    title = doc.add_heading(brand.name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Apply Brand Tone to Formatting
    # (Simulating brand-aware styling)
    if brand.tone.lower() == "professional":
        font_name = "Arial"
        font_size = 11
    elif brand.tone.lower() == "creative":
        font_name = "Georgia"
        font_size = 12
    else:
        font_name = "Calibri"
        font_size = 11

    # 3. Add Content
    for paragraph in content.split("\n\n"):
        p = doc.add_paragraph(paragraph)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)

    # 4. Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Document generated at: {output_path}")
    
    return os.path.abspath(output_path)
